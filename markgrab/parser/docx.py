"""DOCX parser — extract text from Word documents using python-docx."""

import logging
from io import BytesIO

from markgrab.result import ExtractResult
from markgrab.utils import detect_language

logger = logging.getLogger(__name__)


class DocxParser:
    """Extract text from DOCX bytes using python-docx."""

    def parse(self, data: bytes, url: str) -> ExtractResult:
        """Parse DOCX content.

        Args:
            data: Raw DOCX bytes.
            url: Source URL.
        """
        from docx import Document

        doc = Document(BytesIO(data))

        # Extract title from core properties
        title = ""
        if doc.core_properties and doc.core_properties.title:
            title = doc.core_properties.title.strip()

        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        full_text = "\n\n".join(paragraphs)

        # Build markdown (respect heading styles)
        md_lines = []
        if title:
            md_lines.append(f"# {title}\n")
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or "").lower() if para.style else ""
            if "heading 1" in style_name:
                md_lines.append(f"# {text}\n")
            elif "heading 2" in style_name:
                md_lines.append(f"## {text}\n")
            elif "heading 3" in style_name:
                md_lines.append(f"### {text}\n")
            elif "heading" in style_name:
                md_lines.append(f"#### {text}\n")
            elif "list" in style_name:
                md_lines.append(f"- {text}")
            else:
                md_lines.append(text)
                md_lines.append("")
        markdown = "\n".join(md_lines).strip()

        language = detect_language(full_text)

        docx_metadata: dict = {}
        props = doc.core_properties
        if props:
            if props.author:
                docx_metadata["author"] = props.author
            if props.created:
                docx_metadata["created"] = str(props.created)
            if props.modified:
                docx_metadata["modified"] = str(props.modified)
            if props.subject:
                docx_metadata["subject"] = props.subject

        return ExtractResult(
            title=title or "DOCX Document",
            text=full_text,
            markdown=markdown,
            word_count=len(full_text.split()),
            language=language,
            content_type="document",
            source_url=url,
            metadata=docx_metadata,
        )
